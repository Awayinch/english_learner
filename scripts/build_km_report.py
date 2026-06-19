from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "面向雅思备考的个人语言知识管理系统大作业报告.docx"

TITLE = "雅思备考知识管理系统"
FULL_TITLE = "面向个人语言能力构建的垂直知识管理系统开发与实践"
COURSE = "知识工程管理课程大作业"
SCHOOL = "大连理工大学"
TODAY = date.today().strftime("%Y年%m月%d日")


INK = RGBColor(0, 0, 0)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 96, 110)


def set_run_font(run, name="宋体", size=12, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(paragraph, *, first_line=True, after=0, before=0, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Cm(0.74)


def add_para(doc, text="", *, style=None, align=None, first_line=True, size=12, bold=False, color=INK, after=0):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    set_paragraph_format(p, first_line=first_line, after=after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = None
    if level == 1:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(11)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        set_run_font(run, name="黑体", size=15, bold=True, color=INK)
    elif level == 2:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_run_font(run, name="黑体", size=14, bold=True, color=INK)
    else:
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_run_font(run, name="黑体", size=12, bold=True, color=DARK_BLUE)
    return p


def add_center_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(11)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, name="黑体", size=15, bold=True, color=INK)
    return p


def add_bullets(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=12)


def add_numbered(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=12)


def set_cell_text(cell, text, *, bold=False, fill=None, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.text = ""
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color="C9D1D9"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_in: Sequence[float]):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            row.cells[idx].width = Inches(width)
    set_table_borders(table)


def add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float], caption: str):
    cap = add_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=10.5, bold=True, after=4)
    cap.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, fill="F2F4F7", align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], text, align=alignment)
    add_para(doc, "", first_line=False, after=4)
    return table


def add_flow_figure(doc, caption: str, labels: Sequence[tuple[str, str]], fill="E8EEF5"):
    cap = add_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=10.5, bold=True, after=4)
    cap.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(labels))
    table.style = "Table Grid"
    set_table_width(table, [6.5 / len(labels)] * len(labels))
    for idx, (title, desc) in enumerate(labels):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, f"{title}\n{desc}", bold=True, fill=fill, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    add_para(doc, "", first_line=False, after=4)


def add_screenshot_box(doc, caption: str, note: str):
    cap = add_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=10.5, bold=True, after=4)
    cap.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(note)
    set_run_font(run, name="宋体", size=11, bold=True, color=MUTED)
    add_para(doc, "", first_line=False, after=4)


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def paragraph_bottom_border(paragraph, color="B7B7B7", size="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def add_page_number(paragraph, *, roman=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("- ")
    set_run_font(run, size=9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* ROMAN" if roman else "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    tail = paragraph.add_run(" -")
    set_run_font(tail, size=9)


def apply_section_geometry(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(2.5)
    section.footer_distance = Cm(2)


def set_section_header_footer(section, *, roman=False, start=1):
    apply_section_geometry(section)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    set_page_number_start(section, start)
    header = section.header.paragraphs[0]
    clear_paragraph(header)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(TITLE)
    set_run_font(run, size=10.5, color=MUTED)
    paragraph_bottom_border(header)

    footer = section.footer.paragraphs[0]
    clear_paragraph(footer)
    add_page_number(footer, roman=roman)


def setup_document(doc: Document):
    section = doc.sections[0]
    apply_section_geometry(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    set_section_header_footer(section, roman=True, start=1)


def page_break(doc):
    doc.add_page_break()


def add_cover(doc):
    for _ in range(5):
        add_para(doc, "", first_line=False)
    add_para(doc, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=22, bold=True, after=16)
    add_para(doc, f"（{COURSE}）", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=14, bold=True, color=MUTED, after=72)

    fields = [
        ("学 院（系）：", "____________________________"),
        ("专       业：", "____________________________"),
        ("团 队 成 员：", "____________________________"),
        ("任 课 教 师：", "____________________________"),
        ("完 成 日 期：", TODAY),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.25
        label_run = p.add_run(label)
        set_run_font(label_run, size=14, bold=False)
        value_run = p.add_run(value)
        set_run_font(value_run, size=14)

    for _ in range(6):
        add_para(doc, "", first_line=False)
    add_para(doc, SCHOOL, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=15, bold=True, after=2)
    add_para(doc, "Dalian University of Technology", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=11, color=INK)


def add_abstract(doc):
    add_center_title(doc, "摘     要")
    abstract = (
        "外语学习本质上是学习者对语言知识进行获取、加工、存储、检索和内化的过程。传统英语学习工具常常能够提供对话、翻译或词汇查询功能，"
        "但学习痕迹分散在聊天记录、截图和临时笔记中，难以形成可复用、可追踪、可复盘的个人知识资产。本文以个人项目 LingoLeap 为基础，"
        "围绕雅思备考场景，将其改造为一个垂直领域个人知识管理系统。系统通过生词卡片、历史会话、语境自动化增强、知识资产看板、学习反思记录"
        "和 Obsidian 兼容的 Markdown 导出，实现从原始学习数据到结构化知识卡片的转换。理论上，本文采用 DIKW 模型解释单词、释义、语境和应用能力"
        "之间的层级关系，采用 SECI 知识创造螺旋解释隐性语言经验如何外化、组合并重新内化；同时结合词汇习得、自我调节学习、认知负荷和学习分析"
        "相关研究，将系统指标限定为可观测日志和学习者自评，避免把装饰性统计误解释为真实掌握度。实践结果表明，低成本改造可以显著增强英语学习工具"
        "的知识沉淀能力，使其能够支持课程大作业中的系统设计、指标口径说明、成果展示和持续迭代。该项目也说明，个人语言学习系统不仅可以记录学习内容，"
        "还可以管理学习者对自身知识状态和学习策略的认识。"
    )
    add_para(doc, abstract, after=6)
    add_para(doc, "关键词：个人知识管理；雅思备考；知识资产；学习分析；自我调节学习", first_line=False, bold=True)


def add_toc(doc):
    add_center_title(doc, "目     录")
    entries = [
        ("摘 要", "II"),
        ("1 研究背景与问题提出", "1"),
        ("2 理论基础与文献依据", "2"),
        ("3 研究方法与数据来源", "5"),
        ("4 需求分析与系统定位", "7"),
        ("5 系统总体架构设计", "9"),
        ("6 核心功能模块设计", "11"),
        ("7 指标体系与知识资产看板", "15"),
        ("8 系统实现与运行流程", "17"),
        ("9 成果展示与配图说明", "19"),
        ("10 项目管理、备份与迭代", "21"),
        ("11 总结与展望", "23"),
        ("参考文献", "24"),
        ("附录A 截图清单与验收材料", "26"),
    ]
    for title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.25
        left = p.add_run(title)
        set_run_font(left, size=12)
        dots = p.add_run(" " + "." * max(4, 42 - len(title)) + " ")
        set_run_font(dots, size=12)
        right = p.add_run(page)
        set_run_font(right, size=12)


def add_section_page(doc, title: str, paragraphs: Sequence[str], bullets: Sequence[str] | None = None):
    add_heading(doc, title, level=1)
    for text in paragraphs:
        add_para(doc, text)
    if bullets:
        add_bullets(doc, bullets)


def build_report():
    doc = Document()
    setup_document(doc)

    add_cover(doc)
    page_break(doc)
    add_abstract(doc)
    page_break(doc)
    add_toc(doc)
    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_header_footer(body_section, roman=False, start=1)

    add_section_page(
        doc,
        "1 研究背景与问题提出",
        [
            "在数字化学习环境中，语言学习者每天都会接触大量单词、例句、语法现象和对话语境。这些材料在产生时通常只是零散数据，若缺少持续组织，就会沉没在聊天记录和临时文件中。",
            "知识管理课程强调知识的获取、共享、存储和创新。把英语学习系统改造为个人知识管理系统，可以将语言学习从“查一个词、背一个词”的孤立动作，转化为围绕目标能力持续积累知识资产的过程。",
            "本项目选择雅思备考作为垂直场景，是因为雅思学习高度依赖词汇、语境、表达迁移和自我反思。系统不仅要记录词汇本身，还要记录词汇出现在哪里、学习者是否主动使用过、学习后主观负荷如何，以及下一轮复习计划是什么。",
        ],
        [
            "研究对象：基于 React 与 Vite 的 LingoLeap 英语学习项目。",
            "改造目标：面向雅思备考构建个人语言知识管理系统。",
            "交付形态：可运行程序、知识资产看板、Markdown 导出、课程报告与截图证据。",
        ],
    )
    add_flow_figure(
        doc,
        "图1-1 语言学习数据向个人知识资产转化的基本思路",
        [
            ("学习输入", "阅读、对话、文本"),
            ("知识外化", "生词与语境记录"),
            ("知识组合", "标签、双链、增强卡片"),
            ("复习内化", "输出、反思、计划"),
        ],
    )
    page_break(doc)

    add_section_page(
        doc,
        "2 理论基础与文献依据",
        [
            "DIKW 模型将知识处理划分为数据、信息、知识和智慧四个层级。对于语言学习来说，单独的单词只是数据；词性、释义和例句使其成为信息；带标签、来源语境、复习计划和主动使用痕迹的卡片才更接近可复用知识；基于这些记录调整学习策略，则体现为实践层面的智慧。",
            "Nonaka 与 Takeuchi 提出的 SECI 模型强调知识在隐性与显性之间循环转换。语言学习中的语感和表达经验属于隐性知识，记录生词和例句是外化，整合标签、同义词、词根和 Markdown 双链是组合化，复习和再次输出则推动内化。",
            "词汇习得研究中，Nation 强调词汇知识包含 form、meaning、use；Laufer 与 Hulstijn 的投入量假说指出任务需求、搜索和评价会影响词汇保持。系统的语境增强模块正是围绕这些维度补充同义表达、构词提示、词族和例句。",
        ],
    )
    add_table(
        doc,
        ["理论", "系统映射", "在报告中的作用"],
        [
            ["DIKW", "单词、释义、语境、复习计划逐层加工", "解释知识层级"],
            ["SECI", "输入、外化、组合、内化形成循环", "解释知识创造过程"],
            ["词汇知识框架", "词形、词义、词性、语境证据", "定义卡片完整率"],
            ["学习分析", "只统计可观测日志和自评记录", "约束指标口径"],
        ],
        [1.35, 2.7, 2.45],
        "表2-1 理论基础与系统功能的对应关系",
    )
    page_break(doc)

    add_section_page(
        doc,
        "2.1 自我调节学习与认知负荷",
        [
            "自我调节学习理论强调学习者需要经历目标设定、过程监控、策略调整和结果反思。仅记录词汇数量无法说明学习者是否知道自己下一步应该做什么，因此系统增加了“学习反思/自评记录”模块。",
            "认知负荷理论提醒我们，学习任务过难或界面信息过载都会影响学习效果。由于认知负荷和学习焦虑无法从词汇数量直接推断，系统采用 1-5 分自评方式采集主观状态，并将其作为过程性指标，而不是最终成绩指标。",
            "这种设计使报告可以明确区分客观日志与主观自评：词汇卡片、语境暴露和主动使用来自系统日志；认知负荷、焦虑和信心来自学习者填写；主动回忆成功率和间隔复习则仍然需要未来测验数据。",
        ],
    )
    add_flow_figure(
        doc,
        "图2-1 学习反思模块对应的自我调节闭环",
        [
            ("计划", "下一步学习目标"),
            ("监控", "负荷、焦虑、信心"),
            ("调整", "识别待强化节点"),
            ("反思", "导出并复盘"),
        ],
        fill="E6FFFA",
    )
    page_break(doc)

    add_section_page(
        doc,
        "3 研究方法与数据来源",
        [
            "本项目采用设计型研究与轻量学习分析结合的方法。设计型研究关注在真实场景中持续改造系统，并通过迭代验证功能是否服务于学习目标；学习分析关注从学习行为数据中提取可解释指标。",
            "本项目不追求一次性重写系统，而是遵循低风险、可截图、可回退的原则，在现有侧边栏中增加知识资产看板、语境自动化增强、学习反思自评和 Markdown 导出。",
            "数据来源全部来自本地程序可以观察或用户主动输入的记录。系统不会把没有证据的数据包装为结果，也不会根据词汇数量直接声称学习者已经掌握某个词。",
        ],
    )
    add_table(
        doc,
        ["数据类型", "来源", "可支持的指标", "边界"],
        [
            ["词汇卡片", "vocabulary 本地状态", "词汇卡片数、完整率", "不能代表掌握度"],
            ["历史会话", "聊天 session 与消息", "语境暴露、会话深度", "只能说明出现过"],
            ["用户消息", "role 为 user 的文本", "主动使用事件", "不能判断使用质量"],
            ["自评记录", "学习反思模块", "负荷、焦虑、信心", "主观过程指标"],
            ["导出文件", "Markdown 下载", "知识沉淀证据", "需结合笔记软件展示"],
        ],
        [1.15, 1.7, 2.0, 1.65],
        "表3-1 研究数据来源与指标边界",
    )
    page_break(doc)

    add_section_page(
        doc,
        "4 需求分析与系统定位",
        [
            "目标用户是需要在较短时间内准备雅思考试的个人学习者。该类用户常见痛点包括：资料来源复杂，生词和例句无法长期保存；学习过程中缺少对知识存量和薄弱节点的可视化判断；AI 对话产生的内容难以沉淀到 Obsidian 等个人知识库。",
            "因此系统定位不是通用聊天工具，而是面向雅思备考的垂直个人知识管理系统。系统应帮助学习者完成从输入、记录、加工、存储到复习的闭环，并能为课程大作业提供清晰的理论解释和界面证据。",
        ],
        [
            "知识采集：支持手动添加、AI 导入和待查询队列。",
            "知识加工：支持词性、释义、语境、同义词、构词提示和例句增强。",
            "知识盘点：用看板展示可观测指标和待强化节点。",
            "知识沉淀：将结果导出为中文 Markdown，兼容 Obsidian 双链。",
            "自我调节：记录学习反思、自评分数和下一步计划。",
        ],
    )
    add_screenshot_box(doc, "图4-1 系统主界面截图占位", "建议替换为：打开 LingoLeap 后，左侧面板展开“知识资产看板”的截图")
    page_break(doc)

    add_section_page(
        doc,
        "5 系统总体架构设计",
        [
            "系统在原有前端项目上进行轻量改造，核心技术栈为 React、Vite、TypeScript 与浏览器本地存储。架构上可以划分为交互层、学习层、知识加工层、知识存储层和外部知识库层。",
            "交互层负责聊天、词汇面板、看板和自评表单；学习层承载 AI 对话、文本导入和词汇记录；知识加工层负责指标计算、语境增强和知识卡片组装；存储层负责 localStorage、词汇列表和导出文件；外部知识库层则对接 Obsidian 或 GitHub Sync。",
        ],
    )
    add_flow_figure(
        doc,
        "图5-1 LingoLeap PKM 系统总体架构",
        [
            ("交互层", "React 面板"),
            ("学习层", "对话与生词"),
            ("加工层", "指标与增强"),
            ("存储层", "本地状态"),
            ("知识库", "Markdown/Obsidian"),
        ],
    )
    add_table(
        doc,
        ["层级", "职责", "本次新增内容"],
        [
            ["交互层", "呈现面板、按钮、表单与下载入口", "知识资产看板、自评模块"],
            ["加工层", "计算指标并组合词汇卡片", "语境增强、待强化节点"],
            ["存储层", "保存词汇、自评和导出结果", "localStorage 反思记录"],
            ["外部知识库", "长期归档与复盘", "中文 Markdown 双链导出"],
        ],
        [1.2, 2.6, 2.7],
        "表5-1 系统层级职责",
    )
    page_break(doc)

    add_section_page(
        doc,
        "6 核心功能模块设计",
        [
            "本次改造聚焦四个可展示模块：知识资产看板、语境自动化增强、学习反思/自评记录和 PKM Markdown 导出。这些模块共同服务于知识管理流程，而不是单纯增加页面装饰。",
            "知识资产看板将本地词汇、历史消息和用户输出转化为指标；语境增强将单一词条扩展为结构化卡片；反思模块将主观学习状态显性化；导出模块把系统内部记录转为外部知识库文件。",
        ],
    )
    add_table(
        doc,
        ["模块", "输入", "处理", "输出"],
        [
            ["知识资产看板", "词卡、会话、消息", "计算完整率、暴露、主动使用", "知识存量与薄弱节点"],
            ["语境增强", "词汇卡片、历史消息", "匹配同义词、词根、语境", "深加工知识卡片"],
            ["学习反思", "用户 1-5 分自评", "本地保存与均值统计", "自我调节学习证据"],
            ["Markdown 导出", "看板、词卡、自评", "生成 YAML 与双链", "Obsidian 兼容文件"],
        ],
        [1.3, 1.8, 2.0, 1.4],
        "表6-1 核心模块输入输出",
    )
    page_break(doc)

    add_section_page(
        doc,
        "6.1 知识资产看板",
        [
            "知识资产看板的设计原则是“能从日志计算的才展示为已观测指标”。系统统计词汇卡片数、语境暴露、主动使用、深加工卡片、卡片完整率、语境化率、输出迁移率和平均会话深度。",
            "看板同时保留待采集指标区域，明确主动回忆成功率与间隔复习逾期数需要后续测验或复习日志。这样可以避免大作业报告中出现没有数据基础的“掌握度”结论。",
        ],
    )
    add_screenshot_box(doc, "图6-1 知识资产看板截图占位", "建议替换为：显示词汇卡片、语境暴露、主动使用、反思记录和已观测指标的截图")
    page_break(doc)

    add_section_page(
        doc,
        "6.2 语境自动化增强模块",
        [
            "语境自动化增强模块用于体现知识管理中的组合化和深加工。系统对每个词汇卡片补充同义表达、词根/构词提示、词族候选、自动例句和来源语境。",
            "增强逻辑优先使用历史对话中的真实语境。如果没有命中历史语境，则生成 IELTS 场景例句并标注增强方式。报告中应说明，自动例句是辅助加工内容，不应被误写为学习者真实输出。",
            "该模块对应 Laufer 与 Hulstijn 的投入量假说：搜索同义表达、评价词族关系和构造例句都可以提高词汇处理深度。",
        ],
    )
    add_screenshot_box(doc, "图6-2 语境增强词卡截图占位", "建议替换为：点击“增强”后，词卡显示同义、构词、词族、例句和来源语境的截图")
    page_break(doc)

    add_section_page(
        doc,
        "6.3 学习反思/自评记录模块",
        [
            "学习反思/自评记录模块是本次新增的轻量功能。学习者可以在看板中记录认知负荷、学习焦虑和复习信心三个 1-5 分指标，并填写下一步学习计划。",
            "该模块的数据保存在浏览器 localStorage 中，并参与看板统计与 Markdown 导出。与词汇日志不同，自评数据属于主观过程指标，反映学习者对当前状态的监控和计划，而不是考试成绩。",
            "从知识管理角度看，自评记录把学习者脑中的隐性状态外化为显性记录，使系统不仅管理语言知识点，也管理学习者对自身学习过程的认识。",
        ],
    )
    add_screenshot_box(doc, "图6-3 学习反思/自评记录截图占位", "建议替换为：保存一次反思后，显示最近记录与均值统计的截图")
    page_break(doc)

    add_section_page(
        doc,
        "6.4 PKM Markdown 导出模块",
        [
            "PKM Markdown 导出模块把系统内部状态转换为可长期保存的知识库文件。导出内容包含中文 YAML Frontmatter、知识资产看板、指标口径说明、DIKW 映射、SECI 映射、近期学习会话、学习反思记录和词汇知识卡片。",
            "每个词汇卡片使用 Obsidian 双链形式，例如 [[词汇/雅思/impact]]，并带有 #词汇/雅思、#个人知识管理/语言学习 等标签。这样导出的文件既能作为学习资料，也能作为大作业成果截图。",
        ],
    )
    add_screenshot_box(doc, "图6-4 Markdown 导出文件截图占位", "建议替换为：下载的 LingoLeap-雅思个人知识库-日期.md 在编辑器或 Obsidian 中打开的截图")
    page_break(doc)

    add_section_page(
        doc,
        "7 指标体系与知识资产看板",
        [
            "指标体系的核心不是追求数字越多越好，而是确保每个数字都有来源、有解释、有边界。系统将指标分为已观测指标、已采集自评指标和待采集指标三类。",
            "已观测指标来自系统日志，例如词汇卡片数、目标词在历史消息中的出现次数、目标词在用户消息中的出现次数。已采集自评指标来自学习者手动填写。待采集指标则明确标注需要未来数据。",
        ],
    )
    add_table(
        doc,
        ["指标", "类型", "计算/采集方式", "解释边界"],
        [
            ["词汇卡片", "已观测", "vocabulary.length", "知识外化规模"],
            ["语境暴露", "已观测", "目标词命中历史消息", "输入痕迹"],
            ["主动使用", "已观测", "目标词命中用户消息", "输出痕迹"],
            ["认知负荷", "自评采集", "1-5 分滑杆均值", "主观状态"],
            ["主动回忆成功率", "待采集", "测验正确次数/总次数", "未来测验模块"],
        ],
        [1.2, 1.2, 2.3, 1.8],
        "表7-1 指标体系与数据来源",
    )
    page_break(doc)

    add_section_page(
        doc,
        "8 系统实现与运行流程",
        [
            "实现层面，本次改造集中在 VocabularyPanel 组件中完成。该组件本来承担生词本、历史对话和导出功能，因此在其中增加看板和自评记录，能最大限度减少对主聊天流程、模型调用和语音播放等基础功能的影响。",
            "知识资产看板通过 sessions、messages、vocabulary 和 pendingWords 计算指标。反思记录通过 localStorage 保存，避免引入后端数据库。Markdown 导出通过浏览器下载机制生成文件，保证演示时不依赖额外服务。",
        ],
    )
    add_flow_figure(
        doc,
        "图8-1 系统运行流程",
        [
            ("输入", "聊天/文本/手动词汇"),
            ("记录", "词卡与会话"),
            ("增强", "同义/构词/例句"),
            ("自评", "负荷/焦虑/信心"),
            ("导出", "Markdown 知识库"),
        ],
        fill="F4F6F9",
    )
    add_numbered(
        doc,
        [
            "学习者在聊天或阅读过程中记录生词。",
            "系统将词汇加入本地词库，并与历史消息进行命中匹配。",
            "学习者点击“增强”生成深加工卡片。",
            "学习后填写反思自评与下一步计划。",
            "点击 PKM 导出，生成可接入 Obsidian 的 Markdown 文件。",
        ],
    )
    page_break(doc)

    add_section_page(
        doc,
        "9 成果展示与配图说明",
        [
            "成果展示部分建议在最终提交前插入四类真实截图：主界面、知识资产看板、语境增强词卡、学习反思记录、Markdown 导出文件以及 Obsidian 双链效果。",
            "本报告已预留截图框，并提供图题。替换截图时应保留图题和编号，使正文引用保持一致。若时间不足，也可以保留系统架构图、指标表和流程图作为配图，但真实截图更能证明程序已经运行。",
        ],
    )
    add_table(
        doc,
        ["截图", "页面位置", "证明内容"],
        [
            ["系统主界面", "http://localhost:5173/", "项目可运行"],
            ["知识资产看板", "左侧面板", "指标来自本地日志"],
            ["学习反思模块", "看板中部", "自评记录可保存"],
            ["语境增强词卡", "生词本列表", "知识卡片深加工"],
            ["Markdown 文件", "下载文件或 Obsidian", "知识外部化与沉淀"],
        ],
        [1.5, 2.0, 3.0],
        "表9-1 建议截图清单",
    )
    page_break(doc)

    add_section_page(
        doc,
        "10 项目管理、备份与迭代",
        [
            "由于课程截止时间紧，本项目采用分支开发与小步提交策略。先将原始项目保持在可回退状态，再在独立分支中加入课程功能，确保任何一次功能改造出现问题时都能回到上一个可运行版本。",
            "迭代顺序遵循“先恢复基础功能，再加看板，再加增强，再加自评，最后写报告”的路线。这样既保证程序可以演示，也保证报告内容与真实功能一致。",
        ],
        [
            "版本控制：使用 Git 分支 codex/pkm-assignment 保存大作业改造。",
            "验证方式：优先运行 npm run build；TypeScript 单文件转译用于快速发现语法问题。",
            "备份策略：重要改造前使用分支或 stash 保留回退点。",
            "交付策略：程序、文档、报告和截图清单同步更新。",
        ],
    )
    page_break(doc)

    add_section_page(
        doc,
        "10.1 测试与验证",
        [
            "测试阶段重点关注三个方面：页面能否正常构建，新增模块是否破坏主流程，导出文件是否包含课程报告需要的证据。由于本项目存在历史 src 目录重复文件，完整 tsc 检查会暴露部分既有类型问题，因此本次以 Vite build 和关键组件转译作为主要验证。",
            "最新构建结果显示，Vite 生产构建能够完成，说明新增学习反思模块和文档改动没有引入阻断性编译错误。后续若继续完善项目，可以逐步清理 src 重复结构，再恢复严格的全量 typecheck。",
        ],
    )
    add_table(
        doc,
        ["验证项", "结果", "说明"],
        [
            ["diff --check", "通过", "无明显空白或补丁格式问题"],
            ["组件 TS 转译", "通过", "VocabularyPanel 无语法错误"],
            ["npm run build", "通过", "Vite 生产构建完成"],
            ["Playwright 截图", "未完成", "本地 npm 缓存权限限制，报告保留截图占位"],
        ],
        [1.6, 1.2, 3.7],
        "表10-1 验证记录",
    )
    page_break(doc)

    add_section_page(
        doc,
        "11 知识工程管理价值分析",
        [
            "从知识工程管理角度看，本项目的价值在于将个人语言学习中难以管理的碎片化内容转化为结构化对象。词汇卡片是知识单元，语境增强是知识加工，Markdown 导出是知识存储，自评记录是学习过程元知识。",
            "系统的另一价值是强调指标边界。许多学习应用容易把“记录了多少词”包装成“掌握了多少词”。本项目明确区分知识资产规模、学习行为痕迹和真实掌握度，体现了数据治理和指标设计的基本意识。",
            "对于课程大作业而言，该系统可以同时展示理论模型、系统架构、数据口径、功能实现和迭代管理，避免报告停留在纯理论或纯截图层面。",
        ],
    )
    add_flow_figure(
        doc,
        "图11-1 本项目中的知识管理闭环",
        [
            ("采集", "词汇/会话"),
            ("整理", "卡片/标签"),
            ("加工", "语境增强"),
            ("评估", "看板/自评"),
            ("沉淀", "Markdown"),
        ],
        fill="EEF2FF",
    )
    page_break(doc)

    add_section_page(
        doc,
        "12 风险、伦理与隐私说明",
        [
            "个人知识管理系统会处理学习者输入文本、词汇记录和学习反思。虽然本项目当前主要使用浏览器本地存储，但报告仍需说明隐私边界：API Key 不应写入报告或截图，导出的 Markdown 文件在分享前应检查是否包含私人聊天内容。",
            "AI 生成的同义词、词根和例句存在不准确风险，因此系统在导出文件中保留“增强方式”字段，并将无法确认的词根标注为待词典确认。后续版本可以接入可溯源词典 API，提高知识卡片可信度。",
            "学习反思数据属于主观状态记录，不应被用于评价或排名，只应服务于学习者自我监控和复盘。",
        ],
    )
    add_bullets(
        doc,
        [
            "截图前隐藏真实 API Key、Token 和代理密钥。",
            "导出 Markdown 前检查是否包含私人对话。",
            "把自动例句表述为系统生成内容，不写成真实学习输出。",
            "把自评指标表述为主观过程数据，不写成成绩或诊断结论。",
        ],
    )
    page_break(doc)

    add_section_page(
        doc,
        "13 局限与未来工作",
        [
            "当前版本仍属于课程作业导向的 MVP，重点是低成本、可运行和可解释。它还不能完整判断学习者是否掌握某词，也没有实现真正的间隔重复算法和错题日志。",
            "未来可以在三方面继续扩展。第一，增加测验记录，采集主动回忆正确率。第二，为每个词汇增加 lastReviewedAt、nextReviewAt 和 reviewCount，支持间隔复习。第三，接入词典 API 或语料库，为同义词、词根和例句提供更可靠的来源。",
        ],
    )
    add_table(
        doc,
        ["未来功能", "新增数据", "知识管理意义"],
        [
            ["测验模块", "正确/错误记录", "从知识资产走向学习成效评估"],
            ["间隔复习", "复习日期与次数", "支持长期保持"],
            ["外部词典", "可溯源释义与例句", "提升知识卡片可信度"],
            ["Obsidian 自动同步", "文件路径与提交记录", "增强知识库联动"],
        ],
        [1.6, 2.0, 2.9],
        "表13-1 未来迭代方向",
    )
    page_break(doc)

    add_section_page(
        doc,
        "14 结 论",
        [
            "本文基于 LingoLeap 英语学习项目，完成了面向雅思备考的个人语言知识管理系统改造。系统新增知识资产看板、语境自动化增强、学习反思/自评记录和 PKM Markdown 导出，将零散的语言学习数据组织为可解释、可保存、可复盘的知识资产。",
            "理论上，系统以 DIKW 解释从单词数据到语言能力的层级加工，以 SECI 解释隐性语言经验和显性知识卡片之间的循环转换，并结合词汇习得、自我调节学习和学习分析来定义指标口径。",
            "实践上，项目证明在不大规模重构原系统的前提下，也可以通过小模块实现贴合知识管理主题的功能增强。该方案适合作为课程大作业展示，也为后续继续完善个人学习系统提供了基础。",
        ],
    )
    page_break(doc)

    add_center_title(doc, "参 考 文 献")
    references = [
        "Ackoff, R. L. (1989). From data to wisdom. Journal of Applied Systems Analysis, 16, 3-9.",
        "Nonaka, I., & Takeuchi, H. (1995). The Knowledge-Creating Company. Oxford University Press.",
        "Nation, I. S. P. (2001). Learning Vocabulary in Another Language. Cambridge University Press.",
        "Nation, I. S. P. (2007). The four strands. Innovation in Language Learning and Teaching, 1(1), 2-13.",
        "Laufer, B., & Hulstijn, J. (2001). Incidental vocabulary acquisition in a second language: The construct of task-induced involvement. Applied Linguistics, 22(1), 1-26.",
        "Roediger, H. L., & Karpicke, J. D. (2006). Test-enhanced learning. Psychological Science, 17(3), 249-255.",
        "Cepeda, N. J., Pashler, H., Vul, E., Wixted, J. T., & Rohrer, D. (2006). Distributed practice in verbal recall tasks. Psychological Bulletin, 132(3), 354-380.",
        "Sweller, J. (1988). Cognitive load during problem solving. Cognitive Science, 12(2), 257-285.",
        "Zimmerman, B. J. (2000). Attaining self-regulation. In Handbook of Self-Regulation. Academic Press.",
        "Siemens, G., & Long, P. (2011). Penetrating the fog: Analytics in learning and education. EDUCAUSE Review, 46(5), 30-40.",
    ]
    for ref in references:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(ref)
        set_run_font(run, size=10.5)
    page_break(doc)

    add_heading(doc, "附录A 截图清单与验收材料", level=1)
    add_para(doc, "为便于最终提交，建议按下表补齐真实截图。截图插入后，右键更新目录页码，并检查图题编号是否连续。")
    add_table(
        doc,
        ["序号", "截图名称", "操作方法"],
        [
            ["1", "主界面", "运行 npm run dev 后打开 http://localhost:5173/"],
            ["2", "知识资产看板", "展开左侧面板，选择“知识资产看板”"],
            ["3", "学习反思记录", "填写滑杆和计划，点击保存"],
            ["4", "语境增强词卡", "在生词本点击“增强”"],
            ["5", "Markdown 导出", "点击 PKM 下载中文知识库文件"],
            ["6", "Obsidian 双链", "将 Markdown 文件放入 Obsidian 后截图"],
        ],
        [0.7, 1.8, 4.0],
        "表A-1 截图与验收材料清单",
    )
    add_para(doc, "附录说明：本报告正文已包含程序功能、理论映射、指标口径和项目管理说明。若课程要求更偏“知识工程管理”，可在第11章增加组织知识管理或知识工程生命周期的讨论；若课程要求更偏“知识管理”，可突出 SECI 与个人知识库实践。")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    output = build_report()
    print(output)
